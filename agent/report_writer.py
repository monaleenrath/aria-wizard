"""
report_writer.py
----------------
Persists the narrative card to disk:
  - Markdown : flat .md file (easy to paste anywhere)
  - Word     : styled .docx with colored section blocks that mirror
               the briefing-card look (headline banner, accent strips).
"""

from __future__ import annotations

import logging
import os
from datetime import datetime
from pathlib import Path

log = logging.getLogger(__name__)


# Section colour palette (R, G, B) — tuned for a board-room slide deck
PALETTE = {
    "headline":  (0x0F, 0x6C, 0xBD),  # blue
    "summary":   (0x37, 0x47, 0x5A),  # slate
    "kpi":       (0x1F, 0x6F, 0x54),  # green
    "drivers":   (0x6F, 0x42, 0xC1),  # purple
    "anomaly":   (0xC4, 0x31, 0x4B),  # red
    "action":    (0xCA, 0x5B, 0x00),  # amber
    "speaker":   (0x55, 0x55, 0x55),  # neutral grey
}


def write_markdown(markdown: str, output_dir: str, reference_date: str) -> str:
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    fname = f"briefing_{reference_date}.md"
    fpath = os.path.join(output_dir, fname)
    with open(fpath, "w", encoding="utf-8") as fh:
        fh.write(markdown)
    log.info("Wrote markdown: %s", fpath)
    return fpath


def _add_card(doc, icon: str, label: str, body: str, color_key: str):
    """One styled 'card' section inside the Word doc."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    r, g, b = PALETTE[color_key]

    # Section header line — icon + label in colour
    header = doc.add_paragraph()
    run_icon = header.add_run(f"{icon}  ")
    run_icon.font.size = Pt(14)
    run_lbl = header.add_run(label)
    run_lbl.bold = True
    run_lbl.font.size = Pt(13)
    run_lbl.font.color.rgb = RGBColor(r, g, b)

    # Body — render markdown-ish content
    for line in body.splitlines():
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("|"):
            # leave markdown table rows as monospace-ish plain text
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)

    # Thin separator
    sep = doc.add_paragraph("_" * 60)
    for run in sep.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(8)


def write_docx(narrative, output_dir: str, reference_date: str) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor

    Path(output_dir).mkdir(parents=True, exist_ok=True)
    fname = f"briefing_{reference_date}.docx"
    fpath = os.path.join(output_dir, fname)

    doc = Document()

    # ---- Banner ---- #
    banner = doc.add_paragraph()
    run = banner.add_run("⚡  DAILY BOARD BRIEFING")
    run.bold = True
    run.font.size = Pt(11)
    r, g, b = PALETTE["headline"]
    run.font.color.rgb = RGBColor(r, g, b)

    # ---- Headline ---- #
    h = doc.add_heading(narrative.headline, level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*PALETTE["headline"])

    sub = doc.add_paragraph(
        f"Reporting date: {reference_date}  •  Generated {datetime.now():%Y-%m-%d %H:%M}  •  Model: {narrative.model}"
    )
    for run in sub.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")

    # ---- Cards ---- #
    _add_card(doc, "📄", "Executive Summary", narrative.exec_summary, "summary")
    _add_card(doc, "📊", "Key Performance Indicators", narrative.kpi_table_md, "kpi")
    _add_card(doc, "🔎", "What Drove the Move", narrative.drivers_md, "drivers")
    _add_card(doc, "⚠️", "Anomaly & Watch-out", narrative.anomaly, "anomaly")
    _add_card(doc, "➡️", "Recommended Action", narrative.recommended_action, "action")
    _add_card(doc, "🎙️", "Speaker Notes", narrative.speaker_notes, "speaker")

    # ---- Footer ---- #
    footer = doc.add_paragraph(
        "Source: Superstore (Tableau Public dashboard). Comparisons are "
        "calendar-day aligned. Anomalies flagged via z-score ≥ 2.0 on a "
        "90-day trailing window."
    )
    for run in footer.runs:
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(fpath)
    log.info("Wrote docx: %s", fpath)
    return fpath
